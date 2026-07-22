"""Local document parsers for `.md`, `.txt`, `.pdf`, and `.docx` files.

Extracts plain text only - never the original file bytes (those are
never stored in PostgreSQL; only the parsed, approved text is). PDF
parsing uses `pypdf.PdfReader.extract_text()` and DOCX parsing iterates
`python-docx` paragraphs; neither library executes embedded macros,
VBA, JavaScript, or OLE objects, and this module never touches those
embedded-object streams even indirectly - only paragraph/page text is
read. No OCR is performed: a PDF with no extractable text (scanned or
image-only) is rejected with a controlled `UnsupportedDocumentError`
rather than silently returning an empty document. No network access
occurs anywhere in this module.
"""

from __future__ import annotations

import hashlib
from dataclasses import dataclass
from pathlib import Path

from stock_research_core.application.exceptions import UnsupportedDocumentError
from stock_research_core.domain.ai_tutor.enums import KnowledgeSourceType

MAX_DOCUMENT_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20 MB

_EXTENSION_TO_SOURCE_TYPE = {
    ".md": KnowledgeSourceType.LOCAL_MARKDOWN,
    ".txt": KnowledgeSourceType.LOCAL_TEXT,
    ".pdf": KnowledgeSourceType.LOCAL_PDF,
    ".docx": KnowledgeSourceType.LOCAL_DOCX,
}

_DOCX_HEADING_PREFIXES = {
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
    "Heading 5": "#####",
    "Heading 6": "######",
    "Title": "#",
}


@dataclass(frozen=True)
class ParsedDocument:
    """The result of parsing one local file - text only, never raw bytes."""

    text: str
    content_hash: str
    source_type: KnowledgeSourceType
    title: str | None


def _normalize_text(raw_text: str) -> str:
    normalized = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    return normalized.replace("\x00", "")


def _guess_title(text: str) -> str | None:
    for line in text.splitlines():
        stripped = line.strip().lstrip("#").strip()
        if stripped:
            return stripped[:300]
    return None


def _parse_markdown_or_text(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise UnsupportedDocumentError("File is not valid UTF-8 text.") from exc


def _parse_pdf(data: bytes) -> str:
    import io

    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as exc:
        raise UnsupportedDocumentError("File is not a readable PDF.") from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001 - any decrypt failure is a controlled rejection
            raise UnsupportedDocumentError("Encrypted PDFs are not supported.") from exc

    pages_text: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            pages_text.append(f"## Page {page_number}\n\n{page_text}")

    combined = "\n\n".join(pages_text).strip()
    if not combined:
        raise UnsupportedDocumentError(
            "No extractable text was found in this PDF (it may be scanned or image-only; OCR is "
            "not supported)."
        )
    return combined


def _parse_docx(data: bytes) -> str:
    import io

    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - any parse failure is a controlled rejection
        raise UnsupportedDocumentError("File is not a readable DOCX document.") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        prefix = _DOCX_HEADING_PREFIXES.get(paragraph.style.name if paragraph.style else "")
        lines.append(f"{prefix} {text}" if prefix else text)

    combined = "\n\n".join(lines).strip()
    if not combined:
        raise UnsupportedDocumentError("No extractable text was found in this DOCX document.")
    return combined


def parse_local_document(
    file_path: Path, *, max_file_size_bytes: int = MAX_DOCUMENT_FILE_SIZE_BYTES
) -> ParsedDocument:
    """Parse a local `.md`/`.txt`/`.pdf`/`.docx` file into approved plain text.

    Raises `UnsupportedDocumentError` for unsupported extensions,
    oversized files, unreadable/encrypted files, and scanned/image-only
    PDFs with no extractable text.
    """
    if not file_path.is_file():
        raise UnsupportedDocumentError(f"'{file_path}' is not a file.")

    extension = file_path.suffix.lower()
    source_type = _EXTENSION_TO_SOURCE_TYPE.get(extension)
    if source_type is None:
        raise UnsupportedDocumentError(
            f"Unsupported file extension '{extension}'. Supported: .md, .txt, .pdf, .docx"
        )

    file_size = file_path.stat().st_size
    if file_size > max_file_size_bytes:
        raise UnsupportedDocumentError(
            f"File '{file_path.name}' is {file_size} bytes, exceeding the "
            f"{max_file_size_bytes}-byte limit."
        )
    if file_size == 0:
        raise UnsupportedDocumentError(f"File '{file_path.name}' is empty.")

    data = file_path.read_bytes()

    if source_type in (KnowledgeSourceType.LOCAL_MARKDOWN, KnowledgeSourceType.LOCAL_TEXT):
        raw_text = _parse_markdown_or_text(data)
    elif source_type == KnowledgeSourceType.LOCAL_PDF:
        raw_text = _parse_pdf(data)
    else:
        raw_text = _parse_docx(data)

    normalized_text = _normalize_text(raw_text).strip()
    if not normalized_text:
        raise UnsupportedDocumentError(f"File '{file_path.name}' contains no usable text content.")

    content_hash = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()
    title = _guess_title(normalized_text)

    return ParsedDocument(text=normalized_text, content_hash=content_hash, source_type=source_type, title=title)
